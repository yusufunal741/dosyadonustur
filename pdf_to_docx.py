import sys
import fitz
from docx import Document

pdf_path = sys.argv[1]
docx_path = sys.argv[2]

pdf = fitz.open(pdf_path)
doc = Document()

for page in pdf:
    text = page.get_text()

    if text.strip():
        for line in text.splitlines():
            doc.add_paragraph(line)
    else:
        doc.add_paragraph("[Bu sayfada çıkarılabilir metin bulunamadı.]")

doc.save(docx_path)
pdf.close()

print("PDF -> DOCX BAŞARILI!")
print("Çıktı:", docx_path)